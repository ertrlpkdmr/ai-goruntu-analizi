"""Rapor.md dosyasini Word (.docx) formatina donusturur"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import os

doc = Document()

# Sayfa kenar bosluklari
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading stilleri
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 51, 102)

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# Rapor icerigini oku
rapor_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RAPOR.md")
with open(rapor_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h.strip()
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def process_inline(paragraph, text):
    """Bold ve code formatlarini isle"""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            paragraph.add_run(part)

i = 0
in_table = False
table_headers = []
table_rows = []
in_code = False
code_lines = []

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Code block
    if line.startswith('```'):
        if in_code:
            # End code block
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(60, 60, 60)
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            code_lines = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Table
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # Separator line?
        if all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        if not in_table:
            in_table = True
            table_headers = cells
        else:
            table_rows.append(cells)
        i += 1
        # Check if next line is still table
        if i >= len(lines) or '|' not in lines[i] or not lines[i].strip().startswith('|'):
            add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        continue

    # Headings
    if line.startswith('# '):
        p = doc.add_heading(line[2:].strip(), level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    # Horizontal rule
    if line.startswith('---'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run('─' * 60)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Bullet points
    if line.startswith('- ') or line.startswith('  - '):
        indent = line.count('  ', 0, line.index('-'))
        text = line.strip().lstrip('- ')
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        process_inline(p, text)
        if indent > 0:
            p.paragraph_format.left_indent = Cm(1.5 * indent)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s(.+)$', line.strip())
    if m:
        p = doc.add_paragraph(style='List Number')
        p.clear()
        process_inline(p, m.group(2))
        i += 1
        continue

    # Empty line
    if line.strip() == '':
        i += 1
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    process_inline(p, line)
    i += 1

# Kaydet
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RAPOR.docx")
doc.save(output_path)
print(f"Rapor kaydedildi: {output_path}")
